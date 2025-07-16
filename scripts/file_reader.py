# file_reader.py
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any
import logging
logger = logging.getLogger(__name__)
from docx import Document
import PyPDF2
import io
import re

from .config import FILE_TYPES, RELIGIOUS_GROUPS, CHUNK_SIZE, MIN_CHUNK_LENGTH
from .utils import (
    extract_clean_text, chunk_text, classify_file_type, 
    determine_religious_group, create_metadata_row,
    preprocess_textbook, preprocess_policy, extract_teacher_answers, preprocess_combined,
    check_chunk_has_url_content, check_chunk_has_ai_summary,
    extract_source_urls_from_chunk
)

class FileReader:
    """Class to handle reading and processing different file types."""
    
    def __init__(self, config):
        self.config = config
        self.processed_data = []
    
    def read_docx_file(self, filepath: Path) -> str:
        """Read and extract text from a DOCX file, including paragraphs and table cells."""
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = []
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)
            text = '\n'.join(text_parts)
            return text
        except Exception as e:
            logger.warning(f"First attempt failed for {filepath}: {e}")
            
            # Try alternative approach for large files or unusual structure
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                with zipfile.ZipFile(filepath, 'r') as zip_file:
                    xml_content = zip_file.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    # Extract text from all text elements
                    text_elements = root.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    text = ' '.join([elem.text for elem in text_elements if elem.text])
                    # If still no text, try alternative namespace or structure
                    if not text.strip():
                        text_elements = root.findall('.//t')
                        text = ' '.join([elem.text for elem in text_elements if elem.text])
                        if not text.strip():
                            text = ''.join(root.itertext())
                logger.info(f"Successfully read large file {filepath} with direct XML parsing")
                return text
            except Exception as e2:
                logger.error(f"Error reading DOCX file {filepath} (both attempts failed): {e2}")
                return ""
    
    # a single interface for all file types
    ## expandable to other files types pdf, text, etc.
    ### this is reading a single file
    def read_file(self, filepath: Path) -> str:
        """Read and extract text from a DOCX file (only)."""
        return self.read_docx_file(filepath)
    
    # process a single file
    ## this is the main function that will be used to process a single file
    ## it will return a list of metadata dictionaries
    ## the metadata dictionaries will contain the file path, the file type, the religious group, and the content
    ## the content will be the text of the file
    ## the file type will be the type of the file
    ## the religious group will be the religious group of the file
    
    def process_file(self, filepath: Path, chunk_large_files: bool = True) -> List[Dict[str, Any]]:
        """
        Process a single file and return metadata rows.
        
        Args:
            filepath: Path to the file
            chunk_large_files: Whether to chunk large files
        
        Returns:
            List of metadata dictionaries
        """
        logger.info(f"Processing file: {filepath.name}")
        raw_text = self.read_file(filepath)
        
        # Better content validation
        if not raw_text or len(raw_text.strip()) < 10:  # At least 10 characters
            logger.warning(f"Insufficient content extracted from {filepath.name} (length: {len(raw_text)})")
            return []
        
        file_type = classify_file_type(filepath.name, self.config['FILE_TYPES'])
        religious_group = determine_religious_group(filepath, self.config['RELIGIOUS_GROUPS'])
        
        # Initialize flags for all file types
        has_url_content = False
        has_ai_summary = False
        url_contents_dict = {}
        source_urls = ''  # Always define source_urls
        
        # Use custom pre-processing for each file type
        if file_type == 'textbook':
            clean_text = preprocess_textbook(raw_text)
        elif file_type == 'policy':
            clean_text = preprocess_policy(raw_text)
        elif file_type == 'interview':
            clean_text = extract_teacher_answers(raw_text)
        elif file_type == 'combined':
            clean_text, url_contents_dict, has_ai_summary, source_urls = preprocess_combined(
                raw_text, 
                fetch_urls=self.config.get('FETCH_URLS', True),
                max_url_chars=self.config.get('MAX_URL_CHARS', 8000),
                use_ai_fallback=self.config.get('USE_OPENAI_FALLBACK', True),
                openai_model=self.config.get('OPENAI_MODEL', 'gpt-4o-mini'),
                max_ai_chars=self.config.get('MAX_AI_SUMMARY_CHARS', 2000)
            )
        else:
            clean_text = extract_clean_text(raw_text)
            source_urls = ''
        
        # Handle large files by chunking
        if chunk_large_files and len(clean_text) > self.config['CHUNK_SIZE']:
            chunks = chunk_text(clean_text, self.config['CHUNK_SIZE'])
            rows = []
            
            # For combined files, find AI summary ranges before chunking
            ai_summary_ranges = []
            if file_type == 'combined' and has_ai_summary:
                ai_markers = ["--- AI SUMMARY", "[AI-GENERATED SUMMARY FROM KNOWLEDGE BASE]", "[AI-GENERATED SUMMARY FROM LIVE CONTENT]"]
                for marker in ai_markers:
                    start_pos = 0
                    while True:
                        pos = clean_text.find(marker, start_pos)
                        if pos == -1:
                            break
                        # Find the end of this AI summary section
                        end_pos = len(clean_text)
                        for next_marker in ai_markers:
                            next_pos = clean_text.find(next_marker, pos + len(marker))
                            if next_pos != -1 and next_pos < end_pos:
                                end_pos = next_pos
                        ai_summary_ranges.append((pos, end_pos))
                        start_pos = end_pos
            
            # Calculate chunk positions to match the chunking logic
            chunk_positions = []
            start = 0
            while start < len(clean_text):
                end = start + self.config['CHUNK_SIZE']
                if end < len(clean_text):
                    # Try to break at sentence boundary (matching chunk_text logic)
                    chunk = clean_text[start:end]
                    last_period = chunk.rfind('.')
                    if last_period > self.config['CHUNK_SIZE'] * 0.7:
                        end = start + last_period + 1
                chunk_positions.append((start, end))
                start = end - 100  # overlap
            
            for i, chunk in enumerate(chunks):
                if len(chunk) >= self.config['MIN_CHUNK_LENGTH']:
                    # Check if this specific chunk has URL content and AI summaries
                    chunk_has_url_content = False
                    chunk_has_ai_summary = False
                    source_url = ''
                    if file_type == 'combined':
                        chunk_has_url_content = check_chunk_has_url_content(chunk, url_contents_dict)
                        chunk_has_ai_summary = check_chunk_has_ai_summary(chunk)
                        # Extract URLs for any URL content (raw or AI summary)
                        if chunk_has_url_content:
                            source_url = extract_source_urls_from_chunk(chunk, url_contents_dict)
                    row = create_metadata_row(
                        filepath=filepath,
                        content=chunk,
                        file_type=file_type,
                        religious_group=religious_group,
                        chunk_id=i,
                        has_url_content=chunk_has_url_content,
                        has_ai_summary=chunk_has_ai_summary
                    )
                    row['chunk_count'] = len(chunks)
                    row['source_url'] = source_url
                    rows.append(row)
            
            logger.info(f"Chunked {filepath.name} into {len(rows)} chunks")
            return rows
        else:
            # For non-chunked files
            file_has_url_content = False
            file_has_ai_summary = False
            
            if file_type == 'combined':
                file_has_url_content = check_chunk_has_url_content(clean_text, url_contents_dict)
                file_has_ai_summary = check_chunk_has_ai_summary(clean_text)
            
            row = create_metadata_row(
                filepath=filepath,
                content=clean_text,
                file_type=file_type,
                religious_group=religious_group,
                has_url_content=file_has_url_content,
                has_ai_summary=file_has_ai_summary
            )
            # For non-chunked files, also add ai_summary_url if relevant
            source_url = ''
            if file_type == 'combined' and file_has_url_content:
                source_url = extract_source_urls_from_chunk(clean_text, url_contents_dict)
            row['source_url'] = source_url
            return [row]
    
    # process a directory
    ## this is the main function that will be used to process all files in a directory
 
    def process_directory(self, directory: Path, chunk_large_files: bool = True) -> List[Dict[str, Any]]:
        """
        Process all files in a directory.
        
        Args:
            directory: Directory to process
            chunk_large_files: Whether to chunk large files
        
        Returns:
            List of all metadata dictionaries
        """
        all_data = []
        supported_extensions = {'.docx', '.pdf', '.txt'}
        
        processed_files = []
        for filepath in directory.rglob('*'):
            if filepath.is_file() and filepath.suffix.lower() in supported_extensions:
                logger.info(f"Found file: {filepath.name}")
                file_data = self.process_file(filepath, chunk_large_files)
                if file_data:
                    processed_files.append(filepath.name)
                all_data.extend(file_data)
        
        logger.info(f"Successfully processed {len(processed_files)} files: {processed_files}")
        
        logger.info(f"Processed {len(all_data)} records from {directory}")
        
        # Add detailed file type breakdown
        if all_data:
            file_types = {}
            for record in all_data:
                file_type = record.get('file_type', 'unknown')
                file_types[file_type] = file_types.get(file_type, 0) + 1
            logger.info(f"File type breakdown for {directory}: {file_types}")
        
        return all_data
    
    # process all data
    ## this is the main function that will be used to process all data
  
    def process_all_data(self, chunk_large_files: bool = True) -> pd.DataFrame:
        """
        Process all data directories and return a DataFrame.
        
        Args:
            chunk_large_files: Whether to chunk large files
        
        Returns:
            DataFrame with all processed data
        """
        all_data = []
        
        # Process Catholic files
        if self.config['CATHOLIC_DIR'].exists():
            catholic_data = self.process_directory(self.config['CATHOLIC_DIR'], chunk_large_files)
            all_data.extend(catholic_data)
        
        # Process Protestant files
        if self.config['PROTESTANT_DIR'].exists():
            protestant_data = self.process_directory(self.config['PROTESTANT_DIR'], chunk_large_files)
            all_data.extend(protestant_data)
        
        # Process Both files
        if self.config['BOTH_DIR'].exists():
            both_data = self.process_directory(self.config['BOTH_DIR'], chunk_large_files)
            all_data.extend(both_data)
        
        # Create DataFrame
        df = pd.DataFrame(all_data)
        
        if not df.empty:
            logger.info(f"Created DataFrame with {len(df)} records")
            logger.info(f"File types: {df['file_type'].value_counts().to_dict()}")
            logger.info(f"Religious groups: {df['religious_group'].value_counts().to_dict()}")
        else:
            logger.warning("No data was processed")
        
        return df 